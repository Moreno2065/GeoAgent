"""
DocumentParser - 文档解析器
============================
支持解析 PDF、Word (.docx/.doc)、纯文本、Markdown 等文档格式。
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Optional, Dict, Any

from .content_container import FileType, FileContent


class DocumentParser:
    """
    文档解析器

    支持的格式：
    - PDF (.pdf)
    - Word (.docx, .doc)
    - 纯文本 (.txt)
    - Markdown (.md)
    - RTF (.rtf)
    """

    def __init__(self):
        self._pdf_parser: Optional[object] = None
        self._word_parser: Optional[object] = None

    def parse(self, file_path: str) -> FileContent:
        """
        解析文档文件

        Args:
            file_path: 文件路径

        Returns:
            FileContent 对象
        """
        path = Path(file_path)

        if not path.exists():
            return FileContent(
                file_name=path.name,
                file_path=str(path),
                file_type=FileType.from_extension(path.suffix),
                error=f"文件不存在: {file_path}",
            )

        suffix = path.suffix.lower()
        file_size = path.stat().st_size

        try:
            if suffix == ".pdf":
                return self._parse_pdf(path, file_size)
            elif suffix in {".docx", ".doc"}:
                return self._parse_word(path, file_size)
            elif suffix in {".txt", ".md", ".rtf"}:
                return self._parse_text(path, file_size)
            else:
                return FileContent(
                    file_name=path.name,
                    file_path=str(path),
                    file_type=FileType.from_extension(suffix),
                    error=f"不支持的文档格式: {suffix}",
                )
        except Exception as e:
            return FileContent(
                file_name=path.name,
                file_path=str(path),
                file_type=FileType.from_extension(suffix),
                error=f"解析失败: {str(e)}",
            )

    def _parse_pdf(self, path: Path, file_size: int) -> FileContent:
        """解析 PDF 文件"""
        text_content = ""
        page_count = 0

        # 尝试使用 PyMuPDF (fitz)
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(path))
            page_count = len(doc)
            text_parts = []

            for page in doc:
                text_parts.append(page.get_text())

            text_content = "\n\n".join(text_parts)
            doc.close()

        except ImportError:
            # 尝试使用 pdfplumber
            try:
                import pdfplumber

                with pdfplumber.open(str(path)) as pdf:
                    page_count = len(pdf.pages)
                    text_parts = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)

                    text_content = "\n\n".join(text_parts)

            except ImportError:
                # 尝试使用 pypdf2
                try:
                    from PyPDF2 import PdfReader

                    reader = PdfReader(str(path))
                    page_count = len(reader.pages)
                    text_parts = []

                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)

                    text_content = "\n\n".join(text_parts)

                except ImportError:
                    return FileContent(
                        file_name=path.name,
                        file_path=str(path),
                        file_type=FileType.PDF,
                        error="未安装 PDF 解析库，请安装: pip install pymupdf 或 pip install pdfplumber",
                        metadata={"file_size": file_size},
                    )

        # 生成摘要
        summary = self._generate_summary(text_content, page_count)

        return FileContent(
            file_name=path.name,
            file_path=str(path),
            file_type=FileType.PDF,
            text_content=text_content.strip(),
            summary=summary,
            metadata={
                "file_size": file_size,
                "page_count": page_count,
                "char_count": len(text_content),
            },
        )

    def _parse_word(self, path: Path, file_size: int) -> FileContent:
        """解析 Word 文件"""
        text_content = ""

        try:
            from docx import Document

            doc = Document(str(path))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text_content = "\n".join(paragraphs)

            # 提取表格
            tables_text = []
            for i, table in enumerate(doc.tables):
                table_text = f"\n[表格 {i + 1}]\n"
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_text += " | ".join(cells) + "\n"
                tables_text.append(table_text)

            if tables_text:
                text_content += "\n\n" + "\n".join(tables_text)

        except ImportError:
            return FileContent(
                file_name=path.name,
                file_path=str(path),
                file_type=FileType.WORD,
                error="未安装 python-docx，请安装: pip install python-docx",
                metadata={"file_size": file_size},
            )

        # 生成摘要
        summary = self._generate_summary(text_content)

        return FileContent(
            file_name=path.name,
            file_path=str(path),
            file_type=FileType.WORD,
            text_content=text_content.strip(),
            summary=summary,
            metadata={
                "file_size": file_size,
                "paragraph_count": len([p for p in text_content.split('\n') if p.strip()]),
                "char_count": len(text_content),
            },
        )

    def _parse_text(self, path: Path, file_size: int) -> FileContent:
        """解析纯文本或 Markdown 文件"""
        try:
            # 尝试 UTF-8 编码
            with open(path, "r", encoding="utf-8") as f:
                text_content = f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(path, "r", encoding="gbk") as f:
                    text_content = f.read()
            except Exception:
                return FileContent(
                    file_name=path.name,
                    file_path=str(path),
                    file_type=FileType.from_extension(path.suffix),
                    error="无法读取文件，请检查编码格式",
                    metadata={"file_size": file_size},
                )

        # 生成摘要
        summary = self._generate_summary(text_content)

        return FileContent(
            file_name=path.name,
            file_path=str(path),
            file_type=FileType.from_extension(path.suffix),
            text_content=text_content.strip(),
            summary=summary,
            metadata={
                "file_size": file_size,
                "encoding": "utf-8",
                "char_count": len(text_content),
                "line_count": len(text_content.splitlines()),
            },
        )

    def _generate_summary(self, text: str, page_count: Optional[int] = None) -> str:
        """生成文档摘要"""
        if not text:
            return "文档内容为空"

        # 取前200字符作为摘要
        preview = text[:200].strip()

        # 添加页数信息（如果有）
        if page_count:
            return f"[{page_count}页] {preview}..."
        else:
            return f"{preview}..."


def extract_text_from_file(file_path: str) -> str:
    """
    便捷函数：直接从文件提取文本

    Args:
        file_path: 文件路径

    Returns:
        提取的文本内容，如果失败返回空字符串
    """
    parser = DocumentParser()
    result = parser.parse(file_path)
    return result.text_content if result.is_success() else ""


def extract_text_with_summary(file_path: str) -> tuple[str, str]:
    """
    便捷函数：提取文本并返回摘要

    Args:
        file_path: 文件路径

    Returns:
        (文本内容, 摘要) 元组
    """
    parser = DocumentParser()
    result = parser.parse(file_path)
    return result.text_content, result.summary
